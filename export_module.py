"""
Professional Export Module for NeuroMarketing Reports
Supports PDF, DOCX, HTML, and JSON export formats
"""

import io
import json
import base64
from datetime import datetime
from typing import Dict, List, Any, Optional
import pandas as pd
import plotly.graph_objects as go
import plotly.express as px
from plotly.io import to_html, to_image
import streamlit as st

# Import for document generation
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.graphics.shapes import Drawing
    from reportlab.graphics.charts.barcharts import VerticalBarChart
    from reportlab.graphics.charts.piecharts import Pie
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

class ProfessionalExporter:
    """Professional report exporter for NeuroMarketing analysis"""
    
    def __init__(self):
        self.company_name = "NeuroMarketing GPT"
        self.logo_path = None  # Can be set to company logo path
        self.brand_colors = {
            'primary': '#667eea',
            'secondary': '#764ba2',
            'accent': '#f093fb',
            'text': '#333333',
            'background': '#ffffff'
        }
    
    def export_comprehensive_report(self, 
                                  analysis_data: Dict, 
                                  format_type: str = 'pdf',
                                  include_visualizations: bool = True,
                                  include_recommendations: bool = True) -> bytes:
        """Export comprehensive analysis report in specified format"""
        
        if format_type.lower() == 'pdf':
            return self._export_pdf_report(analysis_data, include_visualizations, include_recommendations)
        elif format_type.lower() == 'docx':
            return self._export_docx_report(analysis_data, include_visualizations, include_recommendations)
        elif format_type.lower() == 'html':
            return self._export_html_report(analysis_data, include_visualizations, include_recommendations)
        elif format_type.lower() == 'json':
            return self._export_json_report(analysis_data)
        else:
            raise ValueError(f"Unsupported format: {format_type}")
    
    def _export_pdf_report(self, data: Dict, include_viz: bool, include_recs: bool) -> bytes:
        """Generate professional PDF report"""
        if not REPORTLAB_AVAILABLE:
            st.error("ReportLab not available. Please install: pip install reportlab")
            return b""
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            textColor=colors.HexColor(self.brand_colors['primary'])
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            spaceAfter=12,
            textColor=colors.HexColor(self.brand_colors['secondary'])
        )
        
        # Title page
        story.append(Paragraph(f"{self.company_name} Analysis Report", title_style))
        story.append(Spacer(1, 12))
        story.append(Paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", styles['Normal']))
        story.append(Spacer(1, 24))
        
        # Executive Summary
        story.append(Paragraph("Executive Summary", heading_style))
        
        overall_sentiment = data.get('overall_sentiment', {})
        summary_text = f"""
        This comprehensive analysis reveals a {overall_sentiment.get('classification', 'neutral')} sentiment 
        with a polarity score of {overall_sentiment.get('polarity', 0):.3f} and confidence level of 
        {overall_sentiment.get('confidence', 0):.1%}. The analysis covers emotional, business, 
        psychological, and cultural dimensions to provide actionable insights for marketing optimization.
        """
        story.append(Paragraph(summary_text, styles['Normal']))
        story.append(Spacer(1, 18))
        
        # Key Metrics Table
        story.append(Paragraph("Key Performance Metrics", heading_style))
        
        # Create metrics table
        metrics_data = [['Metric', 'Score', 'Classification']]
        
        emotions = data.get('emotions', {})
        if emotions:
            top_emotion = max(emotions.items(), key=lambda x: x[1])
            metrics_data.append(['Top Emotion', f"{top_emotion[1]:.3f}", top_emotion[0].title()])
        
        business = data.get('business_metrics', {})
        if business:
            purchase_intent = business.get('purchase_intent', 0)
            metrics_data.append(['Purchase Intent', f"{purchase_intent:.3f}", 
                               'High' if purchase_intent > 0.7 else 'Medium' if purchase_intent > 0.4 else 'Low'])
        
        psych = data.get('psychological_profile', {})
        if psych:
            attention = psych.get('attention_grabbing', 0)
            metrics_data.append(['Attention Score', f"{attention:.3f}", 
                               'High' if attention > 0.7 else 'Medium' if attention > 0.4 else 'Low'])
        
        metrics_table = Table(metrics_data)
        metrics_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(self.brand_colors['primary'])),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(metrics_table)
        story.append(Spacer(1, 18))
        
        # Detailed Analysis Sections
        if emotions:
            story.append(Paragraph("Emotional Analysis", heading_style))
            emotion_text = "Top emotional responses detected:\n"
            sorted_emotions = sorted(emotions.items(), key=lambda x: x[1], reverse=True)
            for emotion, score in sorted_emotions[:5]:
                emotion_text += f"• {emotion.title()}: {score:.3f}\n"
            story.append(Paragraph(emotion_text, styles['Normal']))
            story.append(Spacer(1, 12))
        
        if business:
            story.append(Paragraph("Business Impact Analysis", heading_style))
            business_text = "Key business metrics:\n"
            for metric, score in business.items():
                business_text += f"• {metric.replace('_', ' ').title()}: {score:.3f}\n"
            story.append(Paragraph(business_text, styles['Normal']))
            story.append(Spacer(1, 12))
        
        # Recommendations
        if include_recs and data.get('recommendations'):
            story.append(Paragraph("Strategic Recommendations", heading_style))
            recs_text = ""
            for i, rec in enumerate(data['recommendations'], 1):
                recs_text += f"{i}. {rec}\n"
            story.append(Paragraph(recs_text, styles['Normal']))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    
    def _export_docx_report(self, data: Dict, include_viz: bool, include_recs: bool) -> bytes:
        """Generate professional DOCX report"""
        if not PYTHON_DOCX_AVAILABLE:
            st.error("python-docx not available. Please install: pip install python-docx")
            return b""
        
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'{self.company_name} Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        doc.add_paragraph(f'Generated on: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
        doc.add_paragraph('')
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        overall_sentiment = data.get('overall_sentiment', {})
        summary = f"""
        This comprehensive analysis reveals a {overall_sentiment.get('classification', 'neutral')} sentiment 
        with a polarity score of {overall_sentiment.get('polarity', 0):.3f} and confidence level of 
        {overall_sentiment.get('confidence', 0):.1%}. The analysis covers emotional, business, 
        psychological, and cultural dimensions to provide actionable insights for marketing optimization.
        """
        doc.add_paragraph(summary)
        
        # Key Metrics
        doc.add_heading('Key Performance Metrics', level=1)
        
        emotions = data.get('emotions', {})
        if emotions:
            doc.add_heading('Emotional Analysis', level=2)
            sorted_emotions = sorted(emotions.items(), key=lambda x: x[1], reverse=True)
            for emotion, score in sorted_emotions[:5]:
                doc.add_paragraph(f'• {emotion.title()}: {score:.3f}', style='List Bullet')
        
        business = data.get('business_metrics', {})
        if business:
            doc.add_heading('Business Metrics', level=2)
            for metric, score in business.items():
                doc.add_paragraph(f'• {metric.replace("_", " ").title()}: {score:.3f}', style='List Bullet')
        
        psych = data.get('psychological_profile', {})
        if psych:
            doc.add_heading('Psychological Profile', level=2)
            for metric, score in psych.items():
                doc.add_paragraph(f'• {metric.replace("_", " ").title()}: {score:.3f}', style='List Bullet')
        
        # Recommendations
        if include_recs and data.get('recommendations'):
            doc.add_heading('Strategic Recommendations', level=1)
            for i, rec in enumerate(data['recommendations'], 1):
                doc.add_paragraph(f'{i}. {rec}', style='List Number')
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    def _export_html_report(self, data: Dict, include_viz: bool, include_recs: bool) -> bytes:
        """Generate professional HTML report"""
        
        # HTML template with professional styling
        html_template = f"""
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>{self.company_name} Analysis Report</title>
            <style>
                body {{
                    font-family: 'Arial', sans-serif;
                    line-height: 1.6;
                    color: {self.brand_colors['text']};
                    max-width: 1200px;
                    margin: 0 auto;
                    padding: 2rem;
                    background-color: {self.brand_colors['background']};
                }}
                .header {{
                    background: linear-gradient(135deg, {self.brand_colors['primary']} 0%, {self.brand_colors['secondary']} 100%);
                    color: white;
                    padding: 2rem;
                    border-radius: 10px;
                    text-align: center;
                    margin-bottom: 2rem;
                }}
                .section {{
                    background: white;
                    padding: 1.5rem;
                    margin-bottom: 1.5rem;
                    border-radius: 8px;
                    box-shadow: 0 2px 10px rgba(0,0,0,0.1);
                }}
                .metric-grid {{
                    display: grid;
                    grid-template-columns: repeat(auto-fit, minmax(250px, 1fr));
                    gap: 1rem;
                    margin: 1rem 0;
                }}
                .metric-card {{
                    background: #f8f9fa;
                    padding: 1rem;
                    border-radius: 6px;
                    border-left: 4px solid {self.brand_colors['primary']};
                }}
                .recommendation {{
                    background: #e8f5e8;
                    padding: 1rem;
                    margin: 0.5rem 0;
                    border-radius: 6px;
                    border-left: 4px solid #28a745;
                }}
                h1, h2, h3 {{
                    color: {self.brand_colors['primary']};
                }}
                .score {{
                    font-size: 1.5em;
                    font-weight: bold;
                    color: {self.brand_colors['secondary']};
                }}
            </style>
        </head>
        <body>
        """
        
        # Header
        html_template += f"""
            <div class="header">
                <h1>{self.company_name} Analysis Report</h1>
                <p>Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}</p>
            </div>
        """
        
        # Executive Summary
        overall_sentiment = data.get('overall_sentiment', {})
        html_template += f"""
            <div class="section">
                <h2>Executive Summary</h2>
                <p>This comprehensive analysis reveals a <strong>{overall_sentiment.get('classification', 'neutral')}</strong> 
                sentiment with a polarity score of <span class="score">{overall_sentiment.get('polarity', 0):.3f}</span> 
                and confidence level of <span class="score">{overall_sentiment.get('confidence', 0):.1%}</span>.</p>
            </div>
        """
        
        # Key Metrics
        html_template += """
            <div class="section">
                <h2>Key Performance Metrics</h2>
                <div class="metric-grid">
        """
        
        emotions = data.get('emotions', {})
        if emotions:
            top_emotion = max(emotions.items(), key=lambda x: x[1])
            html_template += f"""
                <div class="metric-card">
                    <h3>Top Emotion</h3>
                    <div class="score">{top_emotion[0].title()}</div>
                    <p>Score: {top_emotion[1]:.3f}</p>
                </div>
            """
        
        business = data.get('business_metrics', {})
        if business:
            purchase_intent = business.get('purchase_intent', 0)
            html_template += f"""
                <div class="metric-card">
                    <h3>Purchase Intent</h3>
                    <div class="score">{purchase_intent:.3f}</div>
                    <p>{'High' if purchase_intent > 0.7 else 'Medium' if purchase_intent > 0.4 else 'Low'} potential</p>
                </div>
            """
        
        psych = data.get('psychological_profile', {})
        if psych:
            attention = psych.get('attention_grabbing', 0)
            html_template += f"""
                <div class="metric-card">
                    <h3>Attention Score</h3>
                    <div class="score">{attention:.3f}</div>
                    <p>{'High' if attention > 0.7 else 'Medium' if attention > 0.4 else 'Low'} attention grabbing</p>
                </div>
            """
        
        html_template += "</div></div>"
        
        # Detailed Analysis
        if emotions:
            html_template += """
                <div class="section">
                    <h2>Emotional Analysis</h2>
                    <p>Top emotional responses detected:</p>
                    <ul>
            """
            sorted_emotions = sorted(emotions.items(), key=lambda x: x[1], reverse=True)
            for emotion, score in sorted_emotions[:5]:
                html_template += f"<li><strong>{emotion.title()}</strong>: {score:.3f}</li>"
            html_template += "</ul></div>"
        
        if business:
            html_template += """
                <div class="section">
                    <h2>Business Impact Analysis</h2>
                    <p>Key business metrics:</p>
                    <ul>
            """
            for metric, score in business.items():
                html_template += f"<li><strong>{metric.replace('_', ' ').title()}</strong>: {score:.3f}</li>"
            html_template += "</ul></div>"
        
        # Recommendations
        if include_recs and data.get('recommendations'):
            html_template += """
                <div class="section">
                    <h2>Strategic Recommendations</h2>
            """
            for i, rec in enumerate(data['recommendations'], 1):
                html_template += f'<div class="recommendation">{i}. {rec}</div>'
            html_template += "</div>"
        
        # Footer
        html_template += f"""
            <div class="section" style="text-align: center; background: {self.brand_colors['primary']}; color: white;">
                <p>&copy; 2024 {self.company_name}. Professional NeuroMarketing Analysis Platform.</p>
            </div>
        </body>
        </html>
        """
        
        return html_template.encode('utf-8')
    
    def _export_json_report(self, data: Dict) -> bytes:
        """Export analysis data as structured JSON"""
        
        # Add metadata to the export
        export_data = {
            'metadata': {
                'exported_by': self.company_name,
                'export_timestamp': datetime.now().isoformat(),
                'format_version': '1.0',
                'export_type': 'comprehensive_analysis'
            },
            'analysis_data': data
        }
        
        return json.dumps(export_data, indent=2, default=str).encode('utf-8')
    
    def create_visualization_for_export(self, data: Dict, chart_type: str = 'comprehensive') -> str:
        """Create visualization for inclusion in reports"""
        
        if chart_type == 'emotions':
            emotions = data.get('emotions', {})
            if emotions:
                fig = px.bar(
                    x=list(emotions.keys()),
                    y=list(emotions.values()),
                    title="Emotional Analysis Profile",
                    color=list(emotions.values()),
                    color_continuous_scale='viridis'
                )
                return to_html(fig, include_plotlyjs='inline')
        
        elif chart_type == 'business':
            business = data.get('business_metrics', {})
            if business:
                fig = px.radar(
                    r=list(business.values()),
                    theta=list(business.keys()),
                    title="Business Metrics Radar"
                )
                return to_html(fig, include_plotlyjs='inline')
        
        elif chart_type == 'comprehensive':
            # Create a comprehensive dashboard
            from plotly.subplots import make_subplots
            
            fig = make_subplots(
                rows=2, cols=2,
                subplot_titles=('Emotions', 'Business Metrics', 'Psychological', 'Cultural'),
                specs=[[{"type": "bar"}, {"type": "bar"}],
                       [{"type": "bar"}, {"type": "bar"}]]
            )
            
            # Add emotion data
            emotions = data.get('emotions', {})
            if emotions:
                fig.add_bar(
                    x=list(emotions.keys()),
                    y=list(emotions.values()),
                    name='Emotions',
                    row=1, col=1
                )
            
            # Add business metrics
            business = data.get('business_metrics', {})
            if business:
                fig.add_bar(
                    x=list(business.keys()),
                    y=list(business.values()),
                    name='Business',
                    row=1, col=2
                )
            
            # Add psychological profile
            psych = data.get('psychological_profile', {})
            if psych:
                fig.add_bar(
                    x=list(psych.keys()),
                    y=list(psych.values()),
                    name='Psychological',
                    row=2, col=1
                )
            
            # Add cultural analysis
            cultural = data.get('cultural_analysis', {})
            if cultural:
                fig.add_bar(
                    x=list(cultural.keys()),
                    y=list(cultural.values()),
                    name='Cultural',
                    row=2, col=2
                )
            
            fig.update_layout(height=800, title_text="Comprehensive Analysis Dashboard")
            return to_html(fig, include_plotlyjs='inline')
        
        return ""
    
    def generate_executive_summary(self, data: Dict) -> str:
        """Generate executive summary from analysis data"""
        
        overall = data.get('overall_sentiment', {})
        emotions = data.get('emotions', {})
        business = data.get('business_metrics', {})
        
        summary = f"""
        EXECUTIVE SUMMARY
        
        Sentiment Classification: {overall.get('classification', 'Unknown').title()}
        Overall Polarity: {overall.get('polarity', 0):.3f}
        Confidence Level: {overall.get('confidence', 0):.1%}
        
        KEY FINDINGS:
        """
        
        if emotions:
            top_emotion = max(emotions.items(), key=lambda x: x[1])
            summary += f"• Primary emotional response: {top_emotion[0].title()} ({top_emotion[1]:.3f})\n"
        
        if business:
            purchase_intent = business.get('purchase_intent', 0)
            summary += f"• Purchase intent level: {purchase_intent:.3f} ({'High' if purchase_intent > 0.7 else 'Medium' if purchase_intent > 0.4 else 'Low'})\n"
            
            brand_affinity = business.get('brand_affinity', 0)
            summary += f"• Brand affinity score: {brand_affinity:.3f}\n"
        
        recommendations = data.get('recommendations', [])
        if recommendations:
            summary += "\nTOP RECOMMENDATIONS:\n"
            for i, rec in enumerate(recommendations[:3], 1):
                summary += f"{i}. {rec}\n"
        
        return summary
    
    def create_download_link(self, data: bytes, filename: str, label: str) -> str:
        """Create a download link for the exported data"""
        b64 = base64.b64encode(data).decode()
        return f'<a href="data:application/octet-stream;base64,{b64}" download="{filename}">{label}</a>'

# Streamlit integration functions
def render_export_interface(analysis_data: Dict):
    """Render export interface in Streamlit"""
    st.markdown("## 📥 Export Professional Reports")
    
    exporter = ProfessionalExporter()
    
    col1, col2 = st.columns(2)
    
    with col1:
        export_format = st.selectbox(
            "Export Format:",
            ["PDF", "DOCX", "HTML", "JSON"]
        )
        
        include_visualizations = st.checkbox("Include Visualizations", value=True)
    
    with col2:
        include_recommendations = st.checkbox("Include Recommendations", value=True)
        
        export_filename = st.text_input(
            "Filename:",
            value=f"neuromarketing_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        )
    
    if st.button("🚀 Generate Report", type="primary"):
        with st.spinner(f"Generating {export_format} report..."):
            try:
                report_data = exporter.export_comprehensive_report(
                    analysis_data,
                    format_type=export_format.lower(),
                    include_visualizations=include_visualizations,
                    include_recommendations=include_recommendations
                )
                
                # Determine file extension and MIME type
                extensions = {
                    'pdf': 'pdf',
                    'docx': 'docx',
                    'html': 'html',
                    'json': 'json'
                }
                
                mime_types = {
                    'pdf': 'application/pdf',
                    'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    'html': 'text/html',
                    'json': 'application/json'
                }
                
                file_ext = extensions[export_format.lower()]
                mime_type = mime_types[export_format.lower()]
                full_filename = f"{export_filename}.{file_ext}"
                
                # Create download button
                st.download_button(
                    label=f"⬇️ Download {export_format} Report",
                    data=report_data,
                    file_name=full_filename,
                    mime=mime_type
                )
                
                st.success(f"✅ {export_format} report generated successfully!")
                
                # Preview for HTML
                if export_format.lower() == 'html':
                    with st.expander("📖 Preview HTML Report"):
                        st.components.v1.html(report_data.decode('utf-8'), height=600, scrolling=True)
                
                # Summary for JSON
                if export_format.lower() == 'json':
                    with st.expander("📋 JSON Structure Preview"):
                        preview_data = json.loads(report_data.decode('utf-8'))
                        st.json(preview_data)
                
            except Exception as e:
                st.error(f"❌ Export failed: {str(e)}")
                st.error("Please ensure all required dependencies are installed.")

def create_batch_export(analysis_results: List[Dict], formats: List[str]) -> Dict[str, bytes]:
    """Create batch export of multiple analyses"""
    exporter = ProfessionalExporter()
    exports = {}
    
    for format_type in formats:
        combined_data = {
            'batch_analysis': True,
            'total_analyses': len(analysis_results),
            'batch_timestamp': datetime.now().isoformat(),
            'analyses': analysis_results
        }
        
        filename = f"batch_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        try:
            export_data = exporter.export_comprehensive_report(
                combined_data,
                format_type=format_type.lower(),
                include_visualizations=True,
                include_recommendations=True
            )
            exports[format_type] = export_data
        except Exception as e:
            st.error(f"Failed to export {format_type}: {str(e)}")
    
    return exports